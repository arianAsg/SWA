import streamlit as st
from accounting import (
    add_bank, add_check, add_payment_to_transaction, delete_transaction, get_banks, get_checks, get_payments_by_transaction, init_db, add_transaction, get_all_transactions, finance_summary,
    get_financial_reports, add_party, get_parties, add_sim_card,
    get_sim_cards, migrate_db, update_sim_owner, update_transaction
)
from contract_generator import ContractGenerator
import io
import os
import json
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import jdatetime
from typing import Optional

# ---------- تنظیمات اولیه --------------
init_db()
CONTRACT_TYPES = {
    "قرارداد فروش": "فروش",
    "قرارداد خرید/صلح (با مفاد ویژه)": "خرید"
}

migrate_db()
CONTRACTS_FOLDER = "contracts"
ARCHIVE_FILE = os.path.join(CONTRACTS_FOLDER, "archive.json")
LOGO_FOLDER = "logo"
LOGO_PATH = os.path.join(LOGO_FOLDER, "uploaded_logo.png")

os.makedirs(CONTRACTS_FOLDER, exist_ok=True)
os.makedirs(LOGO_FOLDER, exist_ok=True)

# ----------------- استایل سفارشی UI -----------------
st.set_page_config(layout="wide", page_title="سیستم مدیریت سیم کارت")
st.markdown("""
<style>
    * { direction: rtl; text-align: right; font-family: 'B Nazanin', Tahoma, sans-serif; }
    .stTextInput input, .stTextArea textarea { text-align: right; }
    .stSelectbox select { text-align: right; }
    .stDataFrame { width: 100%; }
    .stAlert { text-align: right; }
    .stMarkdown h1, .stMarkdown h2, .stMarkdown h3 { text-align: right; }
</style>
""", unsafe_allow_html=True)

# -------------- نوار کناری: لوگو و آرشیو --------------
def sidebar_content():
    st.sidebar.header("تنظیمات/امکانات")
    
    uploaded_logo = st.sidebar.file_uploader("بارگذاری لوگو/سربرگ (PNG/JPG)", type=['png', 'jpg', 'jpeg'])
    if uploaded_logo:
        with open(LOGO_PATH, "wb") as f:
            f.write(uploaded_logo.getbuffer())
        st.sidebar.success("لوگو با موفقیت ذخیره شد.")

    show_archive = st.sidebar.checkbox("🗂️ مشاهده آرشیو قراردادها")
    if show_archive:
        st.sidebar.subheader("آرشیو قراردادها")
        if os.path.exists(ARCHIVE_FILE):
            with open(ARCHIVE_FILE, "r", encoding='utf-8') as fa:
                archive = json.load(fa)
            for item in reversed(archive):
                st.sidebar.write(f"{item['type']} | {item['datetime']}")
                file_path = os.path.join(CONTRACTS_FOLDER, item["filename"])
                if os.path.exists(file_path):
                    with open(file_path, "rb") as fx:
                        st.sidebar.download_button(
                            label=f"دانلود [{item['filename']}]",
                            data=fx,
                            file_name=item["filename"],
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=item["filename"])
        else:
            st.sidebar.info("هنوز قراردادی ثبت نشده.")

    menu_options = [
        "تولید قرارداد سیم‌کارت",
        "🧾 حسابداری معاملات",
        "📱 مدیریت سیم کارت‌ها",
        "👥 مدیریت مشتریان/فروشندگان",
        "🏦 مدیریت بانک‌ها"
    ]
    selected_menu = st.sidebar.radio("انتخاب بخش", menu_options)
    
    contract_type = None
    if selected_menu == "تولید قرارداد سیم‌کارت":
        contract_type = st.sidebar.radio("نوع قرارداد را انتخاب کنید", list(CONTRACT_TYPES.keys()))
    
    return selected_menu, contract_type

# ----------------- فرم قرارداد ------------------
def show_contract_form():
    st.header("اطلاعات فروشنده")
    cols = st.columns(2)
    with cols[0]:
        seller_name = st.text_input("نام فروشنده*")
        seller_phone = st.text_input("تلفن فروشنده*")
        seller_address = st.text_input("نشانی فروشنده")
    with cols[1]:
        seller_birth = st.text_input("متولد فروشنده")
        seller_issued = st.text_input("صادره از فروشنده")
        seller_national_id = st.text_input("شماره کد ملی فروشنده*")
        seller_child = st.text_input("فرزند فروشنده")

    st.header("اطلاعات خریدار")
    cols = st.columns(2)
    with cols[0]:
        buyer_name = st.text_input("نام خریدار (متصالح)*")
        buyer_phone = st.text_input("تلفن خریدار*")
        buyer_address = st.text_input("نشانی خریدار")
    with cols[1]:
        buyer_birth = st.text_input("متولد خریدار")
        buyer_issued = st.text_input("صادره از خریدار")
        buyer_national_id = st.text_input("شماره کد ملی خریدار*")
        buyer_child = st.text_input("فرزند خریدار")

    st.header("مشخصات سیم کارت")
    sim_number = st.text_input("شماره سیم کارت*")

    st.header("اطلاعات مالی")
    cols = st.columns(2)
    with cols[0]:
        sale_amount = st.text_input("مبلغ مورد معامله (ریال)*")
        sale_amount_toman = st.text_input("مبلغ مورد معامله (تومان)*")
    with cols[1]:
        payment_date = st.text_input("تاریخ و زمان تحویل سیم کارت")
        invoice_amount = st.text_input("مبلغ صورتحساب پرداخت شده (ریال)")
        invoice_date = st.text_input("تاریخ صورتحساب")

    st.header("اطلاعات واریز (حداکثر ۳ ردیف)")
    payment_methods = []
    for i in range(3):
        cols = st.columns(5)
        with cols[0]: description = st.text_input(f"شرح واریز", key=f"desc_{i}")
        with cols[1]: bank = st.text_input(f"بانک", key=f"bank_{i}")
        with cols[2]: amount = st.text_input(f"مبلغ واریزی", key=f"amount_{i}")
        with cols[3]: method = st.text_input(f"نحوه پرداخت", key=f"method_{i}")
        with cols[4]: notes = st.text_input(f"توضیحات", key=f"paynotes_{i}")
        payment_methods.append((description, bank, amount, method, notes))

    notes = st.text_area("توضیحات اضافی:")

    data = dict(
        seller_name=seller_name, seller_phone=seller_phone, seller_address=seller_address,
        seller_birth=seller_birth, seller_issued=seller_issued, 
        seller_national_id=seller_national_id, seller_child=seller_child,
        buyer_name=buyer_name, buyer_phone=buyer_phone, buyer_address=buyer_address,
        buyer_birth=buyer_birth, buyer_issued=buyer_issued, 
        buyer_national_id=buyer_national_id, buyer_child=buyer_child,
        sim_number=sim_number, sale_amount=sale_amount, sale_amount_toman=sale_amount_toman,
        payment_date=payment_date, invoice_amount=invoice_amount, invoice_date=invoice_date,
        payment_methods=payment_methods, notes=notes
    )
    return data

# ----------------- مدیریت سیم کارت‌ها ------------------
def sim_management_tab():
    st.header("مدیریت سیم کارت‌ها")
    
    tabs = st.tabs(["ثبت سیم کارت جدید", "لیست سیم کارت‌ها", "تغییر مالکیت"])
    
    with tabs[0]:
        with st.form("sim_card_form"):
            cols = st.columns(2)
            with cols[0]:
                number = st.text_input("شماره سیم کارت*")
                operator = st.selectbox("اپراتور*", ["همراه اول", "ایرانسل", "رایتل"])
                purchase_price = st.number_input("قیمت خرید (ریال)", min_value=0)
            with cols[1]:
                purchase_date = st.text_input("تاریخ خرید (اختیاری)")
                owner_name = st.text_input("مالک فعلی (اختیاری)")
                notes = st.text_area("توضیحات")
            
            if st.form_submit_button("ثبت سیم کارت"):
                if number and operator:
                    # یافتن ID مالک اگر وجود دارد
                    owner_id = None
                    if owner_name:
                        parties = get_parties()
                        matching_parties = [p for p in parties if owner_name.lower() in p["name"].lower()]
                        if matching_parties:
                            owner_id = matching_parties[0]["id"]
                    
                    add_sim_card(
                        number=number,
                        operator=operator,
                        purchase_price=purchase_price,
                        purchase_date=purchase_date,
                        current_owner_id=owner_id,
                        notes=notes
                    )
                    st.success("سیم کارت با موفقیت ثبت شد.")
                    st.experimental_rerun()
                else:
                    st.error("پر کردن فیلدهای ستاره‌دار الزامی است.")
    
    with tabs[1]:
        st.subheader("لیست سیم کارت‌ها")
        sim_cards = get_sim_cards()
        if sim_cards:
            df = pd.DataFrame(sim_cards)
            st.dataframe(df)
            
            # دکمه دانلود
            csv = df.to_csv(index=False).encode('utf-8-sig')
            st.download_button(
                label="دانلود لیست سیم کارت‌ها (CSV)",
                data=csv,
                file_name="sim_cards.csv",
                mime="text/csv"
            )
        else:
            st.info("هنوز سیم کارتی ثبت نشده است.")
    
    with tabs[2]:
        st.subheader("تغییر مالکیت سیم کارت")
        sim_cards = get_sim_cards()
        parties = get_parties()
        
        if sim_cards and parties:
            sim_options = [f"{sc['number']} ({sc['operator']})" for sc in sim_cards]
            selected_sim = st.selectbox("انتخاب سیم کارت", sim_options)
            
            party_options = [p["name"] for p in parties]
            selected_party = st.selectbox("انتخاب مالک جدید", [""] + party_options)
            
            sale_price = st.number_input("قیمت فروش (ریال)", min_value=0)
            
            if st.button("ثبت تغییر مالکیت"):
                if selected_sim and selected_party:
                    sim_number = selected_sim.split(" ")[0]
                    sim_id = next((sc["id"] for sc in sim_cards if sc["number"] == sim_number), None)
                    party_id = next((p["id"] for p in parties if p["name"] == selected_party), None)
                    
                    if sim_id and party_id:
                        update_sim_owner(sim_id, party_id, sale_price)
                        st.success("مالکیت سیم کارت با موفقیت به روز شد.")
                        st.experimental_rerun()
                    else:
                        st.error("خطا در یافتن سیم کارت یا مالک")
                else:
                    st.error("لطفاً سیم کارت و مالک جدید را انتخاب کنید")
        else:
            st.warning("برای تغییر مالکیت، حداقل باید یک سیم کارت و یک طرف حساب وجود داشته باشد.")

# ----------------- مدیریت مشتریان/فروشندگان ------------------
def parties_management_tab():
    st.header("مدیریت مشتریان و فروشندگان")

    tabs = st.tabs(["ثبت طرف حساب جدید", "لیست طرف‌های حساب"])

    with tabs[0]:
        with st.form("party_form"):
            cols = st.columns(2)
            with cols[0]:
                name = st.text_input("نام کامل*")
                phone = st.text_input("تلفن ثابت")
                mobile = st.text_input("شماره موبایل*")
                national_id = st.text_input("کد ملی*")
                initial_balance = st.number_input("مانده اولیه حساب (ریال)", min_value=0, step=1000)
            with cols[1]:
                address = st.text_input("آدرس")
                party_type = st.selectbox("نوع طرف حساب*", ["مشتری", "همکار", "سایر"])
                account_status = st.selectbox("وضعیت طرف حساب", ["طلبکار", "بدهکار"])
                notes = st.text_area("توضیحات")

            if st.form_submit_button("ثبت طرف حساب"):
                if name and mobile and national_id:
                    add_party(
                        name=name,
                        phone=phone,
                        mobile=mobile,
                        national_id=national_id,
                        address=address,
                        party_type=party_type,
                        account_status=account_status,
                        initial_balance=initial_balance,
                        notes=notes
                    )
                    st.success("طرف حساب با موفقیت ثبت شد.")
                    st.experimental_rerun()
                else:
                    st.error("پر کردن فیلدهای ستاره‌دار الزامی است.")
    
    with tabs[1]:
        st.subheader("لیست طرف‌های حساب")
        parties = get_parties()
        if parties:
            df = pd.DataFrame(parties)
            st.dataframe(df)
            
            # دکمه دانلود
            csv = df.to_csv(index=False).encode('utf-8-sig')
            st.download_button(
                label="دانلود لیست طرف‌های حساب (CSV)",
                data=csv,
                file_name="parties.csv",
                mime="text/csv"
            )
        else:
            st.info("هنوز طرف حسابی ثبت نشده است.")

# ----------------- حسابداری معاملات ------------------
def accounting_tab():
    st.title("🧾 حسابداری خرید و فروش سیم‌کارت")

    tabs = st.tabs(["داشبورد", "ثبت تراکنش", "لیست تراکنش‌ها", "گزارشات مالی"])

    # ================== 📊 داشبورد ==================
    with tabs[0]:
        st.subheader("خلاصه مالی")
        summary = finance_summary()
        col1, col2, col3 = st.columns(3)
        col1.metric("موجودی کل", f"{summary['balance']:,} ریال")
        col2.metric("کل دریافتی‌ها", f"{summary['total_income']:,} ریال")
        col3.metric("کل پرداختی‌ها", f"{summary['total_outcome']:,} ریال")

        st.subheader("گردش مالی ماهانه")
        reports = get_financial_reports()
        if reports['monthly']:
            df_monthly = pd.DataFrame(reports['monthly'], columns=["ماه", "درآمد", "هزینه", "مانده"])
            st.line_chart(df_monthly.set_index("ماه"))
        else:
            st.info("داده‌ای برای نمایش وجود ندارد")

    # ================== 📝 ثبت تراکنش ==================
    with tabs[1]:
        st.subheader("ثبت تراکنش جدید")

        # فرم اصلی تراکنش
        with st.form("transaction_form"):
            cols = st.columns(2)
            with cols[0]:
                tx_type = st.selectbox("نوع تراکنش*", ["دریافت فروش", "پرداخت خرید", "دریافت وام", "پرداخت وام", "سایر"])
                parties = get_parties()
                party_options = [""] + [f"{p['name']} ({p['type']})" for p in parties]
                selected_party = st.selectbox("طرف حساب", party_options)

                # انتخاب سیم کارت
                sim_cards = get_sim_cards()
                sim_options = [""] + [f"{sc['number']} ({sc['operator']})" for sc in sim_cards]
                selected_sim = st.selectbox("سیم کارت مرتبط", sim_options)

            with cols[1]:
                contract_choices = [""]
                if os.path.exists(ARCHIVE_FILE):
                    with open(ARCHIVE_FILE, "r", encoding='utf-8') as fa:
                        archive_list = json.load(fa)
                        contract_choices += [f['filename'] for f in archive_list]
                contract_file = st.selectbox("قرارداد مرتبط", contract_choices)
                description = st.text_area("توضیحات")

            # پرداخت‌های چندگانه
            if "payment_rows" not in st.session_state:
                st.session_state["payment_rows"] = 1
            if st.form_submit_button("➕ افزودن ردیف پرداخت", help="ابتدا روی این کلیک کن تا سطر جدید اضافه بشه"):
                st.session_state["payment_rows"] += 1

            payments_data = []
            for i in range(st.session_state["payment_rows"]):
                c = st.columns([2, 2, 2, 2, 3])
                method = c[0].selectbox("روش پرداخت", ["نقدی", "کارت به کارت", "حواله بانکی", "چک"], key=f"pmethod_{i}")
                amount_pm = c[1].number_input("مبلغ (ریال)", min_value=0, step=10000, key=f"pamount_{i}")
                bank_acc = c[2].text_input("حساب/کارت", key=f"pbank_{i}")
                ref_num = c[3].text_input("شماره پیگیری", key=f"pref_{i}")
                notes_pm = c[4].text_input("توضیحات", key=f"pnotes_{i}")
                payments_data.append((method, amount_pm, bank_acc, ref_num, notes_pm))

            if st.form_submit_button("ثبت تراکنش"):
                total_amount = sum(p[1] for p in payments_data)
                if total_amount > 0:
                    # پیدا کردن ID طرف حساب
                    party_id = None
                    if selected_party:
                        party_name = selected_party.split(" (")[0]
                        party_id = next((p["id"] for p in parties if p["name"] == party_name), None)
                    # پیدا کردن ID سیم کارت
                    sim_card_id = None
                    if selected_sim:
                        sim_number = selected_sim.split(" ")[0]
                        sim_card_id = next((sc["id"] for sc in sim_cards if sc["number"] == sim_number), None)

                    tx_id = add_transaction(
                        tx_type=tx_type,
                        amount=total_amount,
                        description=description,
                        contract_file=contract_file,
                        party_id=party_id,
                        sim_card_id=sim_card_id
                    )
                    for method, amount_pm, bank_acc, ref_num, notes_pm in payments_data:
                        if amount_pm > 0:
                            add_payment_to_transaction(tx_id, method, amount_pm, bank_acc, ref_num, notes_pm)

                    st.success("تراکنش و پرداخت‌ها ثبت شدند.")
                    st.experimental_rerun()
                else:
                    st.error("مجموع مبالغ پرداخت باید بیشتر از صفر باشد.")

        # فرم سریع طرف حساب (جدا از فرم تراکنش)
        with st.expander("➕ ثبت سریع طرف حساب جدید"):
            with st.form("quick_party_form"):
                quick_name = st.text_input("نام کامل")
                quick_type = st.selectbox("نوع", ["مشتری", "همکار", "سایر"])
                quick_mobile = st.text_input("شماره موبایل")
                quick_national_id = st.text_input("کد ملی")
                if st.form_submit_button("ثبت طرف حساب جدید"):
                    if quick_name and quick_mobile and quick_national_id:
                        add_party(name=quick_name, mobile=quick_mobile,
                                  national_id=quick_national_id, party_type=quick_type)
                        st.success("طرف حساب افزوده شد.")
                        st.experimental_rerun()
                    else:
                        st.error("پر کردن نام، موبایل و کد ملی اجباری است.")

    # ================== 📜 لیست تراکنش‌ها ==================
    with tabs[2]:
        st.subheader("لیست تراکنش‌ها")
        transactions = get_all_transactions()
        if transactions:
            for tx in transactions:
                st.markdown(f"**{tx['id']}** | {tx['tx_type']} | {tx['amount']:,} ریال")
                payments = get_payments_by_transaction(tx['id'])
                if payments:
                    for p in payments:
                        st.write(f"▫ {p['payment_method']}: {p['amount']:,} ریال ({p['bank_account']}) [{p['reference_number']}]")
                cols = st.columns(2)
                if cols[0].button("✏ ویرایش", key=f"edit_{tx['id']}"):
                    st.warning("ویرایش تراکنش هنوز پیاده‌سازی نشده!")
                if cols[1].button("🗑 حذف", key=f"del_{tx['id']}"):
                    delete_transaction(tx['id'])
                    st.warning("تراکنش حذف شد.")
                    st.experimental_rerun()
        else:
            st.info("هیچ تراکنشی ثبت نشده.")

    # ================== 📈 گزارشات مالی ==================
    with tabs[3]:
        st.subheader("گزارشات مالی")
        col1, col2 = st.columns(2)
        with col1:
            start_date = st.date_input("تاریخ شروع", value=None)
        with col2:
            end_date = st.date_input("تاریخ پایان", value=None)
        if st.button("اعمال فیلتر"):
            start_str = start_date.strftime("%Y-%m-%d") if start_date else None
            end_str = end_date.strftime("%Y-%m-%d") if end_date else None
            reports = get_financial_reports(start_str, end_str)
        else:
            reports = get_financial_reports()
        if reports['by_operator']:
            st.subheader("تراکنش‌ها بر اساس اپراتور")
            df_operator = pd.DataFrame(reports['by_operator'], columns=["اپراتور", "تعداد تراکنش", "جمع مبلغ"])
            st.dataframe(df_operator)
            st.bar_chart(df_operator.set_index("اپراتور")["جمع مبلغ"])
        else:
            st.info("تراکنشی مرتبط با سیم کارت‌ها وجود ندارد")
def banks_management_tab():
    st.header("🏦 مدیریت بانک‌ها")

    with st.form("bank_form"):
        name = st.text_input("نام بانک*")
        account_number = st.text_input("شماره حساب/کارت*")
        owner = st.text_input("نام صاحب حساب")
        notes = st.text_area("توضیحات")
        if st.form_submit_button("ثبت بانک"):
            if name and account_number:
                add_bank(name, account_number, owner, notes)
                st.success("بانک ثبت شد.")
                st.experimental_rerun()
            else:
                st.error("نام و شماره حساب اجباری است.")

    st.subheader("لیست بانک‌ها")
    banks = get_banks()
    if banks:
        df = pd.DataFrame(banks)
        st.dataframe(df)
    else:
        st.info("هیچ بانکی ثبت نشده است.")
def checks_management_tab():
    st.header("📑 مدیریت چک‌ها")

    with st.form("check_form"):
        check_number = st.text_input("شماره چک*")
        type_ = st.selectbox("نوع چک", ["دریافت", "پرداخت"])
        banks = get_banks()
        bank_options = [""] + [f"{b['name']} - {b['account_number']}" for b in banks]
        selected_bank = st.selectbox("بانک*", bank_options)
        amount = st.number_input("مبلغ (ریال)*", min_value=0, step=10000)
        due_date = st.date_input("تاریخ سررسید")
        status = st.selectbox("وضعیت", ["در جریان", "وصول شد", "برگشتی"])
        notes = st.text_area("توضیحات")

        if st.form_submit_button("ثبت چک"):
            if check_number and selected_bank and amount > 0:
                bank_id = next((b["id"] for b in banks if f"{b['name']} - {b['account_number']}" == selected_bank), None)
                add_check(check_number, type_, bank_id, amount, due_date.strftime("%Y-%m-%d"), status, notes)
                st.success("چک با موفقیت ثبت شد.")
                st.experimental_rerun()
            else:
                st.error("فیلدهای ستاره‌دار را پر کنید.")

    st.subheader("لیست چک‌ها")
    chs = get_checks()
    if chs:
        df = pd.DataFrame(chs)
        st.dataframe(df)
    else:
        st.info("هیچ چکی ثبت نشده است.")
# ----------------- تولید قرارداد ------------------
def generate_contract(contract_type, contract_data):
    if CONTRACT_TYPES[contract_type] == "فروش":
        generator = ContractGenerator()
        word_file = generator.generate_contract(contract_data)
    else:
        word_file = generate_buy_contract(contract_data)
    return word_file

def generate_buy_contract(contract_data):
    doc = Document()
    if os.path.exists(LOGO_PATH):
        doc.add_picture(LOGO_PATH, width=Pt(100))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def rtl(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)

    rtl("بسمه تعالی", bold=True)
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.RIGHT

    rtl(f"فروشنده: {contract_data['seller_name']}")
    rtl(f"تلفن: {contract_data['seller_phone']}")
    rtl(f"نشانی: {contract_data['seller_address']}")
    rtl(f"متولد: {contract_data['seller_birth']}   صادره از: {contract_data['seller_issued']}   شماره کد ملی: {contract_data['seller_national_id']}   فرزند: {contract_data['seller_child']}")

    rtl(f"متصالح (خریدار): {contract_data['buyer_name']}")
    rtl(f"تلفن: {contract_data['buyer_phone']}")
    rtl(f"نشانی: {contract_data['buyer_address']}")
    rtl(f"متولد: {contract_data['buyer_birth']}   صادره از: {contract_data['buyer_issued']}   شماره کد ملی: {contract_data['buyer_national_id']}   فرزند: {contract_data['buyer_child']}")

    rtl(f"مورد فروش: کلیه حقوق عینه، متصوره و فرضیه متعلق به یک رشته سیم کارت شرکت همراه اول به شماره {contract_data['sim_number']}\n"
        "اعم از حق الامتیاز و حق الاشتراک و وام و ودیعه متعلقه احتمالی به نحوی که دیگر هیچگونه حق و ادعایی برای فروشنده در مورد صلح باقی نماند و خریدار قائم مقام قانونی و رسمی فروشنده در شرکت همراه اول می باشد تا مطابق مقررات بنام و نفع خود استفاده نماید.")
    rtl(f"مبلغ مورد فروش: مبلغ {contract_data['sale_amount']} ریال معادل {contract_data['sale_amount_toman']} تومان که تمامی آن به اقرار تسلیم فروشنده گردیده است.")

    table = doc.add_table(rows=1, cols=5)
    hdrs = ["توضیحات", "مبلغ واریزی (ریال)", "بانک", "شرح واریز", "نحوه پرداخت"]
    for i, h in enumerate(hdrs):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = "B Nazanin"
    for payment in contract_data['payment_methods']:
        if any(payment):
            row = table.add_row().cells
            for i, item in enumerate(payment):
                row[i].text = str(item)
                row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                row[i].paragraphs[0].runs[0].font.name = "B Nazanin"
    rtl(f"تاریخ و زمان تحویل سیم کارت به متصالح: {contract_data['payment_date']}")

    solh_text = (
"""مفاد و شرایط:
1- مورد صلح صحیح و سالم به رویت متصالح رسیده و متصالح اقرار به دریافت و تصرف صحیح و سالم آن نموده است.
2- هزینه کلیه مکالمات داخل و خارج کشور تا زمان تنظیم صلحنامه به عهده متصالح خواهد بود.
3- متصالح متعهد به همکاری و حضور در تمام مراجع قانونی و قضایی در صورت لزوم می‌باشد.
4- مسئولیت کامل هرگونه سوءاستفاده یا مزاحمت و پرداخت حقوق و دیون مربوطه از زمان تنظیم صلحنامه به عهده متصالح است.
5- متصالح ضامن کشف فساد احتمالی گردید و تعهد به جبران خسارت دارد.
6- سیم کارت تلفن همراه به صورت مال الاجاره می‌باشد و هزینه‌های ما به التفاوت به عهده متصالح است.
7- متصالح هیچگونه حقی نسبت به قطع و سلب امتیاز نخواهد داشت.
8- در صورت کشف فساد مبلغ سیم کارت به خریدار عودت خواهد شد.
9- این سیم کارت به صورت اسقاط کافه خیارات حتی خیار غبن تنظیم و بر اساس مواد 10، 190 و 362 قانون مدنی معتبر است.
""")
    rtl(solh_text)
    rtl(f"توضیحات: {contract_data['notes']}")
    rtl("شاهد                                     شاهد         خریدار         فروشنده")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def save_contract_file(word_file, contract_type):
    now_jalali = jdatetime.datetime.now().strftime("%Y-%m-%d_%H%M%S")
    filename = f"contract_{contract_type}_{now_jalali}.docx"
    file_path = os.path.join(CONTRACTS_FOLDER, filename)
    with open(file_path, "wb") as f:
        f.write(word_file.getbuffer())
    
    try:
        with open(ARCHIVE_FILE, "r", encoding='utf-8') as fa:
            archive = json.load(fa)
    except:
        archive = []
    
    archive.append({
        "type": contract_type,
        "filename": filename,
        "datetime": now_jalali
    })
    
    with open(ARCHIVE_FILE, "w", encoding='utf-8') as fa:
        json.dump(archive, fa, ensure_ascii=False, indent=2)
    
    return file_path

# ----------------- اجرای اصلی برنامه ------------------
def main():
    selected_menu, contract_type = sidebar_content()
    
    if selected_menu == "تولید قرارداد سیم‌کارت":
        st.header(f"تولید قرارداد {contract_type}")
        contract_data = show_contract_form()
        if st.button("📝 تولید و ذخیره قرارداد"):
            if not contract_type or contract_type not in CONTRACT_TYPES:
                st.error("لطفاً نوع قرارداد را انتخاب کنید.")
            else:
                word_file = generate_contract(contract_type, contract_data)
                file_path = save_contract_file(word_file, CONTRACT_TYPES[contract_type])
                filename = os.path.basename(file_path)
                st.success(f"فایل قرارداد با نام {filename} ثبت و ذخیره شد.")
                st.download_button(
                    "⬇️ دانلود فایل Word همین قرارداد", 
                    data=word_file, 
                    file_name=filename
                )
                
                # ثبت خودکار تراکنش مالی
                try:
                    amount = int(contract_data['sale_amount'].replace(',', ''))
                    tx_type = "دریافت فروش" if CONTRACT_TYPES[contract_type] == "فروش" else "پرداخت خرید"
                    party_name = contract_data['buyer_name'] if CONTRACT_TYPES[contract_type] == "فروش" else contract_data['seller_name']
                    
                    add_transaction(
                        tx_type=tx_type,
                        amount=amount,
                        description=f"قرارداد {CONTRACT_TYPES[contract_type]} سیم کارت {contract_data['sim_number']}",
                        contract_file=filename,
                        party_name=party_name
                    )
                    st.info("تراکنش مالی مرتبط نیز به صورت خودکار ثبت شد.")
                except:
                    st.warning("ثبت خودکار تراکنش مالی با خطا مواجه شد. لطفاً به صورت دستی ثبت کنید.")
    
    elif selected_menu == "🧾 حسابداری معاملات":
        accounting_tab()
    
    elif selected_menu == "📱 مدیریت سیم کارت‌ها":
        sim_management_tab()
    
    elif selected_menu == "👥 مدیریت مشتریان/فروشندگان":
        parties_management_tab()

if __name__ == "__main__":
    main()