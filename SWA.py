import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import io

# تزریق CSS برای راست‌چین کردن تمام عناصر
st.markdown("""
<style>
    * {
        direction: rtl;
        text-align: right;
        font-family: 'B Nazanin', Tahoma, sans-serif;
    }
    .stTextInput input, .stTextArea textarea {
        text-align: right;
    }
    .stSelectbox select {
        text-align: right;
    }
</style>
""", unsafe_allow_html=True)

def add_rtl_paragraph(doc, text, bold=False):
    """تابع کمکی برای اضافه کردن پاراگراف راست‌چین"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    return paragraph

def generate_word_contract():
    st.title("📝 تولید متن قرارداد فروش سیم کارت")
    
    # دریافت اطلاعات از کاربر
    st.header("🔵 اطلاعات فروشنده")
    col1, col2, col3 = st.columns(3)
    with col1:
        seller_birth = st.text_input("متولد (فروشنده):", key="seller_birth")
        seller_phone = st.text_input("تلفن (فروشنده):", key="seller_phone")
    with col2:
        seller_issued = st.text_input("صادره از (فروشنده):", key="seller_issued")
        seller_address = st.text_input("نشانی (فروشنده):", key="seller_address")
    with col3:
        seller_national_id = st.text_input("شماره کد ملی (فروشنده):", key="seller_national_id")
        seller_name = st.text_input("نام فروشنده:", key="seller_name")
    seller_child = st.text_input("فرزند (فروشنده):", key="seller_child")
    
    st.header("🔵 اطلاعات خریدار")
    col1, col2, col3 = st.columns(3)
    with col1:
        buyer_birth = st.text_input("متولد (خریدار):", key="buyer_birth")
        buyer_phone = st.text_input("تلفن (خریدار):", key="buyer_phone")
    with col2:
        buyer_issued = st.text_input("صادره از (خریدار):", key="buyer_issued")
        buyer_address = st.text_input("نشانی (خریدار):", key="buyer_address")
    with col3:
        buyer_national_id = st.text_input("شماره کد ملی (خریدار):", key="buyer_national_id")
        buyer_name = st.text_input("نام خریدار:", key="buyer_name")
    buyer_child = st.text_input("فرزند (خریدار):", key="buyer_child")
    
    st.header("📱 اطلاعات سیم کارت")
    sim_number = st.text_input("شماره سیم کارت:", key="sim_number")
    col1, col2 = st.columns(2)
    with col1:
        sale_amount = st.text_input("مبلغ مورد فروش (ریال):", key="sale_amount")
    with col2:
        sale_amount_toman = st.text_input("مبلغ مورد فروش (تومان):", key="sale_amount_toman")
    
    st.header("💵 اطلاعات پرداخت")
    payment_date = st.text_input("تاریخ و زمان تحویل سیم کارت:", key="payment_date")
    col1, col2 = st.columns(2)
    with col1:
        invoice_amount = st.text_input("مبلغ صورتحساب پرداخت شده (ریال):", key="invoice_amount")
    with col2:
        invoice_date = st.text_input("تاریخ صورتحساب و آبونمان:", key="invoice_date")
    
    st.header("🏦 اطلاعات واریز")
    st.write("لطفاً اطلاعات واریز را وارد کنید (حداکثر 3 ردیف):")
    payment_methods = []
    for i in range(3):
        st.subheader(f"ردیف {i+1}")
        cols = st.columns(5)
        with cols[0]:
            desc = st.text_input(f"شرح واریز", key=f"desc_{i}")
        with cols[1]:
            bank = st.text_input(f"بانک", key=f"bank_{i}")
        with cols[2]:
            amount = st.text_input(f"مبلغ واریزی (ریال)", key=f"amount_{i}")
        with cols[3]:
            method = st.text_input(f"نحوه پرداخت", key=f"method_{i}")
        with cols[4]:
            notes = st.text_input(f"توضیحات", key=f"notes_{i}")
        payment_methods.append((desc, bank, amount, method, notes))
    
    notes = st.text_area("توضیحات اضافی:", key="notes")
    
    if st.button("🖨️ تولید فایل Word", type="primary"):
        # ایجاد سند Word
        doc = Document()
        
        # تنظیم فونت برای متن فارسی
        style = doc.styles['Normal']
        font = style.font
        font.name = 'B Nazanin'
        font.size = Pt(12)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # اضافه کردن عنوان
        add_rtl_paragraph(doc, "بسمه تعالی", bold=True)
        
        # اضافه کردن فاصله
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # اطلاعات فروشنده و خریدار
        seller_info = f"متولد:\t{seller_birth}\t\tصادره از:\t{seller_issued}\t\tشماره کد ملی:\t{seller_national_id}\t\tفرزند:\t{seller_child}\t\tفروشنده:\t{seller_name}"
        add_rtl_paragraph(doc, seller_info)
        
        seller_contact = f"تلفن:\t{seller_phone}\t\tنشانی:\t{seller_address}"
        add_rtl_paragraph(doc, seller_contact)
        
        buyer_info = f"متولد:\t{buyer_birth}\t\tصادره از:\t{buyer_issued}\t\tشماره کد ملی:\t{buyer_national_id}\t\tفرزند:\t{buyer_child}\t\tخریدار:\t{buyer_name}"
        add_rtl_paragraph(doc, buyer_info)
        
        buyer_contact = f"تلفن:\t{buyer_phone}\t\tنشانی:\t{buyer_address}"
        add_rtl_paragraph(doc, buyer_contact)
        
        # متن اصلی قرارداد
        contract_text = f"""
و شماره {sim_number} مورد فروش: کلیه حقوق عینه، متصوره و فرضیه متعلق به یک رشته سیم کارت شرکت همراه اول به شماره
اعم از حق الامتیاز و حق الاشتراک و وام و ودیعه متعلقه احتمالی به نحویکه دیگر هیچگونه حق و ادعایی برای فروشنده در مورد سیم کارت
فروش باقی نماند و خریدار قائم مقام قانونی و رسمی فروشنده در شرکت همراه اول می باشد تا مطابق مقررات بنام و نفع خود استفاده نماید.
تومان که تماماً به اقراره تسلیم فروشنده گردیده است. {sale_amount_toman} ریال معادل {sale_amount} مبلغ مورد فروش: مبلغ
"""
        add_rtl_paragraph(doc, contract_text)
        
        # جدول واریزها
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # هدرهای جدول
        hdr_cells = table.rows[0].cells
        headers = ['توضیحات', 'مبلغ واریزی (ریال)', 'بانک', 'شرح واریز', 'نحوه پرداخت']
        
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[i].paragraphs[0].runs[0].font.name = 'B Nazanin'
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # اضافه کردن داده‌های واریز
        for payment in payment_methods:
            if any(payment):
                row_cells = table.add_row().cells
                for i, item in enumerate(payment):
                    row_cells[i].text = str(item)
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    row_cells[i].paragraphs[0].runs[0].font.name = 'B Nazanin'
        
        # ادامه متن قرارداد
        remaining_text = f"""
با توجه به ماده 390 قانون مدنی در خصوص ضمان درک از آنجایی که فروشنده مبیع مورد معامله را به استناد اسناد صدوری از مخابرات و در ید بودن سیم کارت در زمان انتقال به خود هیچ اطلاعی از معاملات قبلی قبل از مالک شدن خودش ندارد و به دلیل جلوگیری از ضمان قانونی فروشنده مبنی بر مستحق الغیر بودن مبی در یده‌ای قبل فروشنده ضمن انتقال رسمی خط مورد معامله درج عدم ضمان نسبت به خریدار را در قرارداد مزبور و درج و خریدار نیز درکمال آگاهی و صحت عقل این عدم ضمان را می‌پذیرد.

می باشد. مورخ: تاریخ و زمان تحویل سیم کارت به مصالح ساعت: {payment_date}
ریال توسط فروشنده پرداخت شده است. {invoice_amount} به مبلغ: صورتحساب و آبونمان خط مذکور تا تاریخ: {invoice_date}
صحیح و سالم به رویت خریدار رسیده و خریدار اقرار به دریافت و تصرف صحیح و سالم آن نموده است. مورد فروش در تاریخ {payment_date}
می باشد. این سیم کارت به صورت

توضیحات: {notes}


شاهد                                     شاهد         خریدار         فروشنده
"""
        add_rtl_paragraph(doc, remaining_text)
        
        # ذخیره فایل در حافظه
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # ایجاد دکمه دانلود
        st.download_button(
            label="⬇️ دانلود فایل Word",
            data=file_stream,
            file_name="قرارداد_فروش_سیم_کارت.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        st.success("✅ فایل Word با موفقیت تولید شد!")

if __name__ == "__main__":
    generate_word_contract()