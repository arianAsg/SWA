from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import io

class ContractGenerator:
    def __init__(self):
        self.doc = Document()
        self._setup_document_style()

    def _setup_document_style(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'B Nazanin'
        font.size = Pt(12)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _add_rtl_paragraph(self, text, bold=False):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.space_after = Pt(0)
        return paragraph

    def _add_payment_table(self, payment_methods):
        table = self.doc.add_table(rows=1, cols=5)
        table.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        hdr_cells = table.rows[0].cells
        headers = ['توضیحات', 'مبلغ واریزی (ریال)', 'بانک', 'شرح واریز', 'نحوه پرداخت']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr_cells[i].paragraphs[0].runs[0].font.name = 'B Nazanin'
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        for payment in payment_methods:
            if any(payment):
                row_cells = table.add_row().cells
                for i, item in enumerate(payment):
                    row_cells[i].text = str(item)
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    row_cells[i].paragraphs[0].runs[0].font.name = 'B Nazanin'

    def generate_contract(self, data):
        self._add_rtl_paragraph("بسمه تعالی", bold=True)
        self.doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.RIGHT

        seller_info = (
            f"متولد:\t{data['seller_birth']}\t\tصادره از:\t{data['seller_issued']}\t\t"
            f"شماره کد ملی:\t{data['seller_national_id']}\t\tفرزند:\t{data['seller_child']}\t\t"
            f"فروشنده:\t{data['seller_name']}"
        )
        self._add_rtl_paragraph(seller_info)
        seller_contact = f"تلفن:\t{data['seller_phone']}\t\tنشانی:\t{data['seller_address']}"
        self._add_rtl_paragraph(seller_contact)
        buyer_info = (
            f"متولد:\t{data['buyer_birth']}\t\tصادره از:\t{data['buyer_issued']}\t\t"
            f"شماره کد ملی:\t{data['buyer_national_id']}\t\tفرزند:\t{data['buyer_child']}\t\t"
            f"خریدار:\t{data['buyer_name']}"
        )
        self._add_rtl_paragraph(buyer_info)
        buyer_contact = f"تلفن:\t{data['buyer_phone']}\t\tنشانی:\t{data['buyer_address']}"
        self._add_rtl_paragraph(buyer_contact)

        contract_text = (
            f"\nو شماره {data['sim_number']} مورد فروش: کلیه حقوق عینه، متصوره و فرضیه متعلق به یک رشته سیم کارت شرکت همراه اول به شماره\n"
            f"اعم از حق الامتیاز و حق الاشتراک و وام و ودیعه متعلقه احتمالی به نحویکه دیگر هیچگونه حق و ادعایی برای فروشنده در مورد سیم کارت\n"
            f"فروش باقی نماند و خریدار قائم مقام قانونی و رسمی فروشنده در شرکت همراه اول می باشد تا مطابق مقررات بنام و نفع خود استفاده نماید.\n"
            f"تومان که تماماً به اقراره تسلیم فروشنده گردیده است. {data['sale_amount_toman']} ریال معادل {data['sale_amount']} مبلغ مورد فروش: مبلغ\n"
        )
        self._add_rtl_paragraph(contract_text)
        self._add_payment_table(data['payment_methods'])

        remaining_text = (
            f"\nبا توجه به ماده 390 قانون مدنی در خصوص ضمان درک از آنجایی که فروشنده مبیع مورد معامله را به استناد اسناد صدوری از مخابرات و در ید بودن سیم کارت در زمان انتقال به خود هیچ اطلاعی از معاملات قبلی قبل از مالک شدن خودش ندارد و ..."
            f"می باشد. مورخ: تاریخ و زمان تحویل سیم کارت به مصالح ساعت: {data['payment_date']}\n"
            f"ریال توسط فروشنده پرداخت شده است. {data['invoice_amount']} به مبلغ: صورتحساب و آبونمان خط مذکور تا تاریخ: {data['invoice_date']}\n"
            f"صحیح و سالم به رویت خریدار رسیده و خریدار اقرار به دریافت و تصرف صحیح و سالم آن نموده است. مورد فروش در تاریخ {data['payment_date']}\n"
            f"می باشد. این سیم کارت به صورت\n\n"
            f"توضیحات: {data['notes']}\n\n"
            f"شاهد                                     شاهد         خریدار         فروشنده"
        )
        self._add_rtl_paragraph(remaining_text)

        file_stream = io.BytesIO()
        self.doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
